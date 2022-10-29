#!/usr/bin/env python
# coding: utf-8

# In[17]:


from urllib.request import Request,urlopen
import lxml
from bs4 import BeautifulSoup
import requests
import pandas as pd
import numpy as np
import re
from PIL import Image,ImageDraw,ImageFont
import moviepy.editor as mymovie
from moviepy.editor import *
import docx

user_input = input('enter required product link:-  ')

req = Request(user_input, headers={'User-Agent': 'Mozilla/5.0'})
webpage = urlopen(req).read()

with requests.Session() as c:
    soup = BeautifulSoup(webpage, 'html.parser')
    a=soup.find("div",attrs={"class":"single-product-zoom"})
    n=a.find("img") 
    img_url = n.get('src')
    logo_link = soup.find('span', attrs={'class':'logo'}).find('img').get('src')
    prod_name = n.get('alt')
    reg_price = soup.find('span',id='regular-price').text
    trim = re.compile(r'[^\d.,]+')
    result_reg_price = trim.sub('',reg_price)
    #print(result)
    sale_price = soup.find('span',id='sale-price').text
    trim = re.compile(r'[^\d.,]+')
    result_sale_price = trim.sub('',sale_price)
    #print(result)
    try:
        image_link = requests.get(img_url)
        logo_url = requests.get(logo_link)
        f = open(prod_name + '.jpg','wb')
        l = open(r'Yoshop_logo.jpg','wb')
        f.write(image_link.content)
        l.write(logo_url.content)
        f.close()
        l.close
        print('image saved successfully!!')
    except Exception as e:
        print('image does not exist',e)
        
    prod_img = Image.open(prod_name + '.jpg')
    prod_img = prod_img.resize((320,400))
    logo = Image.open(r'Yoshop_logo.jpg')
    logo = logo.resize((480,100))
    #print(logo_link) 
#     imgs_link=[]
#     item = soup.find('ul',attrs={'class':'single-product-thumbs'})
#     sec_img_link = item.find_all('li',attrs={'class':""})
#     for i in range(len(sec_img_link)):
#         imgs_link.append(sec_img_link[i].get('data-url'))
    
# lk=[]
# cnt=[]
# vid_images=[]
# for i in range(len(imgs_link)):
#     lk.append(requests.get(imgs_link[i]))
#     f = open(prod_name + str(i)+'.jpg','wb')
#     cnt.append(f.write(lk[i].content))
#     vid_images.append(prod_name + str(i)+'.jpg')

bg=Image.open('cr_bg.jpg')
bg = bg.resize((650,630))
bg.paste(prod_img,(16,160))
bg.paste(logo,(75,16))
draw = ImageDraw.Draw(bg)
point1 = 5,128
font1 = ImageFont.truetype('arial.ttf',16)
font2 = ImageFont.truetype('arial.ttf',21)
draw.text(point1,prod_name,'black',font=font1)
point2 = 360,275
draw.text(point2,'Regular Price:- Rs.'+result_reg_price+'\-','black',font = font2)
point3 = 360,320
draw.text(point3,'Sale Price:- Rs.'+result_sale_price+'\-','black',font = font2)
point4 = 26,580
draw.text(point4,'visit us at: www.yoshops.com','black',font=font2)
bg.show()
bg.save('C:\\Users\\Murtuza pipulyawala\\Desktop\\vid_img\\'+prod_name+'.jpg')

path = 'C:\\Users\\Murtuza pipulyawala\\Desktop\\vid_img\\'
out_path = 'C:\\Users\\Murtuza pipulyawala\\Desktop\\img to vid\\'
output_video_name = 'outvid00.mp4'
output_video_path = out_path+output_video_name

pre_imgs = os.listdir(path)
img = []
for i in pre_imgs:
    i = path+i
    img.append(i)

cv2_fourcc = cv2.VideoWriter_fourcc(*'mp4v')

frame = cv2.imread(img[0])
size = list(frame.shape)
del size[2]
size.reverse()

video = cv2.VideoWriter(output_video_path,cv2_fourcc,1,size) #output video,fourcc,fps,size
for i in range(len(img)):
    video.write(cv2.imread(img[i]))
video.release() 

inputvideo = r'C:\Users\Murtuza pipulyawala\Desktop\img to vid\outvid00.mp4'
inputaudio = 'bg_music.mp3'
outputvideo = 'C:\\Users\\Murtuza pipulyawala\\Desktop\\img to vid\\outvid_bgm00.mp4'

videoclip = mymovie.VideoFileClip(inputvideo)
audioclip = mymovie.AudioFileClip(inputaudio)
final_clip = videoclip.set_audio(audioclip)
final_clip.write_videofile(outputvideo,fps=24,codec="libx264",threads=1)


data_science_link = 'https://yoshops.com/products/data-science-basic-training-program-for-everyone-age-10-t0-60'
req = Request(data_science_link, headers={'User-Agent': 'Mozilla/5.0'})
webpage = urlopen(req).read()
with requests.Session() as c:
    soup = BeautifulSoup(webpage, 'html.parser')
    para = soup.find_all("p")
    a=soup.find("div",attrs={"class":"single-product-zoom"})
    n=a.find("img")
    ds_img_url = n.get('src')
    paragr = []
    for i in range(len(para)):
        paragr.append(para[i].get_text())
final_text1 = paragr[29].replace("\r\n"," ").replace("\xa0"," ")
final_text2 = paragr[32].replace("\r\n"," ").replace("\xa0"," ")
ds_img_link = requests.get(ds_img_url) 
f = open('Data Science.jpg','wb')
f.write(ds_img_link.content)

doc = docx.Document()
doc.add_picture('Data Science.jpg',width = docx.shared.Inches(4),height=docx.shared.Inches(3))
doc.add_paragraph(final_text1)
doc.add_paragraph(final_text2)
doc.add_picture('Data Science.jpg',width = docx.shared.Inches(4),height=docx.shared.Inches(3))
doc.save('C:\\Users\\Murtuza pipulyawala\\Desktop\\tsk4.docx')

